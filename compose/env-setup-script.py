import os
import shutil
from faker import Faker
from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from datetime import datetime
import random

fake = Faker()

root_dir = 'test_folder'

structure = {
    "Admin": ["Exécutif", "Gestion"],
    "Finance": ["Fournisseurs", "Clients", "Audits", "Budgétisation", "Paie"],
    "RH": ["Employés/Actifs", "Employés/Inactifs", "Recrutement", "Formation/Intégration", "Formation/Développement", "Politiques"],
    "IT": ["Actifs", "Sauvegardes", "Projets/Projet1", "Projets/Projet2", "Logiciels", "Support"],
    "Legal": ["Contrats", "Conformité", "Litiges"],
    "Marketing": ["Campagnes", "Contenu", "Étude_de_marché", "Réseaux sociaux"],
    "Opérations": ["Logistique", "Production", "Contrôle_qualité", "Planification"],
    "Ventes": ["Propositions", "Rapports", "Territoires"],
    "Partagés": ["Politiques", "Modèles", "Formulaires", "Relations_publiques", "Steve", "Abigail", "John"]
}

topic_words = {
    "Admin": ["executif", "gestion"],
    "Finance": ["fournisseur", "client", "audit", "budgétisation", "paie"],
    "RH": ["actif", "inactif", "recrutement", "formation", "développement", "politique"],
    "IT": ["actif", "sauvegarde", "projet1", "projet2", "logiciel", "support"],
    "Legal": ["contrat", "conformité", "litige"],
    "Marketing": ["campagne", "contenu", "étude", "réseaux"],
    "Opérations": ["logistique", "production", "contrôle_qualité", "planification"],
    "Ventes": ["proposition", "rapport", "territoire"],
    "Partagés": ["politique", "modèle", "formulaire", "relations_publiques", "Steve", "Abigail", "John"]
}

def generate_file_name(department, subfolder, date_suffix, file_type):
    words = topic_words[department]
    if subfolder == "":
        name_prefix = department + "_" + fake.word()
    else:
        name_prefix = words[subfolder] + "_" + fake.word()
    if file_type == 'docx':
        name_prefix += "_document"
    elif file_type == "xlsx":
        name_prefix += "_tableau"
    elif file_type == "pdf":
        name_prefix += "_rapport"
    if date_suffix:
        return f"{name_prefix}_{date_suffix}"
    else:
        return name_prefix

def create_word_doc(path, name):
    doc = Document()
    doc.add_heading(name, 0)
    for _ in range(5):
        doc.add_paragraph(fake.paragraph())
    doc.save(path)

def create_excel_file(path, name):
    wb = Workbook()
    ws = wb.active
    ws.title = "Données"
    ws.append([name])
    for _ in range(10):
        ws.append([fake.word() for _ in range(5)])
    wb.save(path)

def create_pdf_file(path, name):
    c = canvas.Canvas(path, pagesize=letter)
    c.drawString(100, 750, name)
    for i in range(5):
        text = fake.paragraph()
        c.drawString(100, 700 - i * 50, text)
    c.save()

def create_files_in_directory(directory, department, subfolder, num_files):
    existing_files = []
    for i in range(num_files):
        file_type = fake.random_element(elements=('docx', 'xlsx', 'pdf'))
        date_suffix = fake.date_this_decade().strftime('%Y') if fake.boolean(chance_of_getting_true=50) else ''
        if fake.boolean(chance_of_getting_true=10):
            filename = f"temp.{file_type}"
        else:
            filename = f"{generate_file_name(department, subfolder, date_suffix, file_type)}.{file_type}"
        file_path = os.path.join(directory, filename)
        if not os.path.exists(file_path) :
            if file_type == 'docx':
                create_word_doc(file_path, filename)
            elif file_type == 'xlsx':
                create_excel_file(file_path, filename)
            elif file_type == 'pdf':
                create_pdf_file(file_path, filename)
            existing_files.append(file_path)
        else:
            i = i - 1

    num_duplicates = random.randint(1, num_files // 2)
    for _ in range(num_duplicates):
        file_to_duplicate = random.choice(existing_files)
        new_file_path = file_to_duplicate.replace('.', '_copie.')
        shutil.copy(file_to_duplicate, new_file_path)

def create_structure(base_path, structure, num_files):
    for department, subfolders in structure.items():
        department_path = os.path.join(base_path, department)
        os.makedirs(department_path, exist_ok=True)
        create_files_in_directory(department_path, department, "", num_files)
        
        idx = 0
        for subfolder in subfolders:
            subfolder_path = os.path.join(department_path, subfolder)
            os.makedirs(subfolder_path, exist_ok=True)
            create_files_in_directory(subfolder_path, department, idx, num_files)
            idx += 1

os.makedirs(root_dir, exist_ok=True)

num_files = int(input("Enter the number of files to create in each directory: "))

create_structure(root_dir, structure, num_files)

print(f"File structure created under the '{root_dir}' directory.")